"""
generate_tor.py
Task 2 – LKMM-TH TOR Generator

Copies 'Format TOR Pergerakan .docx' and fills ALL content:
  - Paragraph placeholders (with run-merge technique)
  - Table 0: Header (Nama Kegiatan, Pemateri, Jenis)
  - Table 1: Indikator Keberhasilan (3 rows)
  - Table 2: Susunan Acara (3 sesi)
  - Table 3: Kepanitiaan (placeholder NIM/name)
  - Table 4: RAB (Belanja Barang items, totals)
  - Table 5: Signature block

Output: docs/TOR_Kegiatan_LKMM_TH.docx
"""

import os
import shutil

import docx

SOURCE = "docs/Format TOR Pergerakan .docx"
DEST = "docs/TOR_Kegiatan_LKMM_TH.docx"


# ─── Paragraph replacements ──────────────────────────────────────────────────

PARA_REPLACEMENTS = {
    "Nama Kegiatan": "LKMM-TH 2026",
    "-Cantumkan nama kegiatan, tema, serta tagline kegiatan yang diusung jika ada-": (
        "Latihan Kepemimpinan Manajemen Mahasiswa Tingkat Himpunan (LKMM-TH) 2026.\n"
        'Tema: "Sinergi Kepemimpinan Adaptif di Era Digital".\n'
        'Tagline: "Dari Materi ke Aksi, Dalam Satu Hari".'
    ),
    "-Jelaskan bentuk kegiatan dan teknis yang akan dilaksanakan-": (
        "Kegiatan dilaksanakan dalam format One-Day Bootcamp & Mentorship, yaitu satu "
        "hari penuh yang dibagi menjadi tiga sesi:\n"
        "1. Sesi Pagi (Teori Praktis): materi kepemimpinan, komunikasi, dan manajemen "
        "organisasi berbasis studi kasus interaktif.\n"
        "2. Sesi Siang (Sprint FGD): peserta dibagi kelompok dan didampingi mentor "
        "(Kating/Demisioner) untuk menyusun draf Program Kerja secara langsung.\n"
        "3. Sesi Sore (Sidang Mini): presentasi draf Proker, feedback, dan simulasi "
        "sidang dari mentor."
    ),
    "-Sebutkan peserta dan penerima manfaat dari kegiatan tersebut, baik dari jurusan/prodi maupun angkatan-": (
        "Seluruh mahasiswa baru Jurusan Teknik Komputer dan Informatika, "
        "Politeknik Negeri Bandung, Angkatan 2026."
    ),
    "-Sebutkan manfaat dari kegiatan tersebut-": (
        "1. Membekali mahasiswa angkatan 2026 dengan nilai kepemimpinan berintegritas, "
        "kolaboratif, dan manajerial dasar dalam berorganisasi.\n"
        "2. Menghilangkan jeda waktu tanpa pendampingan sehingga pemahaman materi tetap "
        "terjaga dan Proker yang dihasilkan lebih optimal.\n"
        "3. Melatih kemampuan berpikir kritis dan presentasi melalui sidang mini."
    ),
    "nama kegiatan": "LKMM-TH 2026",
}


# ─── Core helpers ─────────────────────────────────────────────────────────────


def merge_and_replace(paragraph, reps):
    """Merge all runs into one, then replace. Handles cross-run placeholders."""
    if not paragraph.runs:
        return
    full_text = "".join(r.text for r in paragraph.runs)
    for old, new in reps.items():
        full_text = full_text.replace(old, new)
    paragraph.runs[0].text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


def set_cell(table, row, col, text):
    """Set a table cell's text. Preserves existing formatting of paragraph runs."""
    cell = table.rows[row].cells[col]
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        cell.text = text


def fill_table_0_header(table):
    """Table 0: Nama Kegiatan, Pemateri, Jenis Kegiatan."""
    set_cell(table, 0, 0, "Nama Kegiatan")
    set_cell(
        table,
        0,
        2,
        "LKMM-TH (Latihan Kepemimpinan Manajemen Mahasiswa Tingkat Himpunan) 2026",
    )
    set_cell(table, 1, 2, "(Disesuaikan dengan pemateri yang ditunjuk oleh jurusan)")
    set_cell(table, 2, 2, "Pelatihan Kepemimpinan Tingkat Himpunan")


def fill_table_1_indikator(table):
    """Table 1: Indikator Keberhasilan Aktivitas (3 data rows)."""
    data = [
        ("Kehadiran peserta", ">= 80% dari total undangan", "90%"),
        ("Tersusunnya draf Proker", "Tiap kelompok menghasilkan 1 draf Proker", "100%"),
        ("Kepuasan peserta", "Survei post-event rata-rata >= 4 dari 5", ">= 4/5"),
    ]
    for i, (indikator, acuan, target) in enumerate(data):
        row_idx = i + 1  # row 0 is header
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 0, indikator)
            set_cell(table, row_idx, 1, acuan)
            set_cell(table, row_idx, 2, target)


def fill_table_2_susunan_acara(table):
    """Table 2: Susunan Acara (schedule on the day)."""
    schedule = [
        (
            "07.30 - 08.00",
            "Registrasi Peserta",
            "Panitia",
            "Ruang Serbaguna JTK",
            "Koordinator Acara",
        ),
        (
            "08.00 - 08.30",
            "Pembukaan & Sambutan",
            "Ketua Himakom / Kajur",
            "Ruang Serbaguna JTK",
            "Koordinator Acara",
        ),
        (
            "08.30 - 10.00",
            "Materi: Kepemimpinan & Komunikasi",
            "Pemateri",
            "Ruang Serbaguna JTK",
            "Koordinator Humas",
        ),
        (
            "10.00 - 10.15",
            "Coffee Break",
            "Panitia",
            "Ruang Serbaguna JTK",
            "Koordinator Konsumsi",
        ),
        (
            "10.15 - 12.00",
            "Materi: Manajemen Organisasi",
            "Pemateri",
            "Ruang Serbaguna JTK",
            "Koordinator Humas",
        ),
        ("12.00 - 13.00", "ISHOMA", "-", "-", "Koordinator Konsumsi"),
        (
            "13.00 - 16.00",
            "FGD Sprint: Penyusunan Proker",
            "Mentor (Kating)",
            "Ruang Kelas JTK",
            "Koordinator FGD",
        ),
        (
            "16.00 - 17.00",
            "Sidang Mini & Presentasi Proker",
            "Mentor & Pemateri",
            "Ruang Serbaguna JTK",
            "Koordinator Acara",
        ),
        (
            "17.00 - 17.30",
            "Penutupan & Evaluasi",
            "Ketua Pelaksana",
            "Ruang Serbaguna JTK",
            "Koordinator Acara",
        ),
    ]
    for i, (waktu, item, pengisi, tempat, pj) in enumerate(schedule):
        row_idx = i + 1
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 0, waktu)
            set_cell(table, row_idx, 1, item)
            set_cell(table, row_idx, 2, pengisi)
            set_cell(table, row_idx, 3, tempat)
            set_cell(table, row_idx, 4, pj)
        else:
            # Add a new row if needed
            row_cells = table.add_row().cells
            row_cells[0].text = waktu
            row_cells[1].text = item
            row_cells[2].text = pengisi
            row_cells[3].text = tempat
            row_cells[4].text = pj


def fill_table_3_kepanitiaan(table):
    """Table 3: Kepanitiaan."""
    # Row 0 is the PJ (pre-filled with Dinanda), leave it.
    # Rows 1-4 are placeholders
    data = [
        # (role already in col0, name, NIM)
        (1, "(Nama Ketua Pelaksana)", "(NIM)"),
        (2, "(Nama Wakil Ketua Pelaksana)", "(NIM)"),
        (3, "(Nama Koordinator)", "(NIM)"),
        (4, "(Nama Anggota)", "(NIM)"),
    ]
    for row_idx, name, nim in data:
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 2, name)
            set_cell(table, row_idx, 3, nim)


def fill_table_4_rab(table):
    """Table 4: RAB (21 rows, 10 cols). Fill header rows + Belanja Barang items."""
    # Row 1: Program Kerja name
    set_cell(table, 1, 0, "Program Kerja LKMM-TH 2026")
    # Row 2: Nama Ormawa
    set_cell(table, 2, 0, "Himpunan Mahasiswa Komputer (HIMAKOM)")
    # Row 4: Tanggal
    set_cell(table, 4, 0, "Dilaksanakan Tanggal : (disesuaikan)")

    # Belanja Barang items (rows 9, 10 under "525112 Belanja Barang")
    # Row 8 = header "525112 Belanja Barang"
    # Rows 9, 10 = empty data rows, row 11 = Sub Total
    # Columns: 0-3=komponen, 4=satuan, 5=volume, 6=harga satuan, 7=polban, 8=sponsor, 9=swadaya
    belanja_barang = [
        # (komponen, satuan, volume, harga_satuan, swadaya)
        (
            "Konsumsi Pemateri & Panitia (makan siang)",
            "Paket",
            "1",
            "Rp 400.000",
            "Rp 400.000",
        ),
        (
            "Print Sertifikat, Kertas FGD, Fotokopi",
            "Paket",
            "1",
            "Rp 150.000",
            "Rp 150.000",
        ),
    ]

    for i, (komponen, satuan, volume, harga, total) in enumerate(belanja_barang):
        row_idx = 9 + i
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 0, komponen)
            set_cell(table, row_idx, 4, satuan)
            set_cell(table, row_idx, 5, volume)
            set_cell(table, row_idx, 6, harga)
            set_cell(table, row_idx, 9, total)  # Swadaya column

    # Sub Total Belanja Barang (row 11)
    set_cell(table, 11, 9, "Rp 550.000")

    # Belanja Jasa (rows 13, 14 under "525113 Belanja Jasa")
    belanja_jasa = [
        ("Honorarium Pemateri (2 sesi)", "Orang", "2", "Rp 0", "Rp 0"),
    ]
    for i, (komponen, satuan, volume, harga, total) in enumerate(belanja_jasa):
        row_idx = 13 + i
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 0, komponen)
            set_cell(table, row_idx, 4, satuan)
            set_cell(table, row_idx, 5, volume)
            set_cell(table, row_idx, 6, harga)
            set_cell(table, row_idx, 9, total)

    # Sub Total Belanja Jasa (row 15)
    set_cell(table, 15, 9, "Rp 0")

    # Belanja Transportasi (rows 17, 18 under "525115 Belanja Transportasi")
    belanja_transport = [
        ("Banner / Spanduk acara (2x3m)", "Lembar", "1", "Rp 150.000", "Rp 150.000"),
        (
            "ID Card Panitia & Biaya Tak Terduga",
            "Paket",
            "1",
            "Rp 100.000",
            "Rp 100.000",
        ),
    ]
    for i, (komponen, satuan, volume, harga, total) in enumerate(belanja_transport):
        row_idx = 17 + i
        if row_idx < len(table.rows):
            set_cell(table, row_idx, 0, komponen)
            set_cell(table, row_idx, 4, satuan)
            set_cell(table, row_idx, 5, volume)
            set_cell(table, row_idx, 6, harga)
            set_cell(table, row_idx, 9, total)

    # Sub Total Belanja Transportasi (row 19)
    set_cell(table, 19, 9, "Rp 250.000")

    # TOTAL (row 20)
    set_cell(table, 20, 9, "Rp 800.000")


def fill_table_5_signature(table):
    """Table 5: Signature block — fill Ketua Pelaksana name placeholder."""
    cell = table.rows[0].cells[2]  # right column
    for para in cell.paragraphs:
        merge_and_replace(
            para,
            {
                "Nama Terang": "(Nama Ketua Pelaksana)",
                "NIM": "(NIM Ketua Pelaksana)",
            },
        )


# ─── Main ─────────────────────────────────────────────────────────────────────


def main():
    if not os.path.exists(SOURCE):
        raise FileNotFoundError(f"Source template not found: {SOURCE}")

    print(f"[1/5] Copying template: {SOURCE} -> {DEST}")
    shutil.copy(SOURCE, DEST)

    print("[2/5] Opening document and applying paragraph replacements ...")
    doc = docx.Document(DEST)

    # Replace paragraph text
    for para in doc.paragraphs:
        merge_and_replace(para, PARA_REPLACEMENTS)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    merge_and_replace(para, PARA_REPLACEMENTS)

    print("[3/5] Filling all tables ...")
    tables = doc.tables

    fill_table_0_header(tables[0])
    fill_table_1_indikator(tables[1])
    fill_table_2_susunan_acara(tables[2])
    fill_table_3_kepanitiaan(tables[3])
    fill_table_4_rab(tables[4])
    fill_table_5_signature(tables[5])

    print("[4/5] Saving document ...")
    doc.save(DEST)

    print(f"\nTOR berhasil dibuat:")
    print(f"  DOCX: {DEST}")


if __name__ == "__main__":
    main()
